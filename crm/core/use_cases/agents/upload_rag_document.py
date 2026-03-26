import io
import re
from dataclasses import dataclass
from uuid import UUID

from core.domain.rag_document import RagDocument
from core.ports.outbound.agent_config_port import AgentConfigPort
from core.ports.outbound.rag_document_port import RagDocumentPort
from core.ports.outbound.tenant_repository import TenantRepositoryPort

_DEFAULT_CHUNK_SIZE = 500
_CHUNK_OVERLAP = 50


def _chunk_text(text: str, chunk_size: int = _DEFAULT_CHUNK_SIZE, overlap: int = _CHUNK_OVERLAP) -> list[str]:
    """Splits text into overlapping chunks, preferring sentence boundaries."""
    text = re.sub(r"\s+", " ", text).strip()
    if not text:
        return []

    if len(text) <= chunk_size:
        return [text]

    chunks: list[str] = []
    start = 0
    while start < len(text):
        end = min(start + chunk_size, len(text))
        # Try to break at sentence boundary
        if end < len(text):
            for sep in (". ", "! ", "? ", "\n", " "):
                pos = text.rfind(sep, start, end)
                if pos > start + chunk_size // 2:
                    end = pos + len(sep)
                    break
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        if end >= len(text):
            break
        start = end - overlap

    return chunks


def _extract_text(filename: str, content: bytes) -> str:
    """Extracts plain text from supported file types."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""

    if ext in ("txt", "md", "rst"):
        return content.decode("utf-8", errors="replace")

    if ext in ("docx", "doc"):
        from docx import Document
        doc = Document(io.BytesIO(content))
        return "\n".join(p.text for p in doc.paragraphs if p.text.strip())

    if ext == "pdf":
        try:
            import fitz  # PyMuPDF
            pdf = fitz.open(stream=content, filetype="pdf")
            return "\n".join(page.get_text() for page in pdf)
        except ImportError:
            raise ValueError("PDF ingestion requires PyMuPDF (pip install pymupdf)")

    raise ValueError(f"Unsupported file type: .{ext}. Supported: txt, md, docx, pdf")


@dataclass(frozen=True)
class UploadRagDocumentRequest:
    tenant_id: UUID
    filename: str
    content: bytes
    doc_name: str = ""  # optional custom name; defaults to filename without extension


class UploadRagDocumentUseCase:

    def __init__(
        self,
        tenant_repo: TenantRepositoryPort,
        config_port: AgentConfigPort,
        rag_port: RagDocumentPort,
    ) -> None:
        self._tenant_repo = tenant_repo
        self._config_port = config_port
        self._rag_port = rag_port

    async def execute(self, request: UploadRagDocumentRequest) -> RagDocument:
        tenant = await self._tenant_repo.get_by_id(request.tenant_id)
        if not tenant:
            raise ValueError("Tenant not found")

        config = self._config_port.read(tenant.agent_id)
        collection = config.get("memory", {}).get(
            "qdrant_rag_collection", f"{tenant.agent_id}_rules"
        )

        name = request.doc_name or request.filename.rsplit(".", 1)[0]
        text = _extract_text(request.filename, request.content)
        chunks = _chunk_text(text)

        if not chunks:
            raise ValueError("No text content extracted from the document")

        count = await self._rag_port.ingest_document(collection, name, chunks)
        return RagDocument(name=name, collection=collection, chunk_count=count)
