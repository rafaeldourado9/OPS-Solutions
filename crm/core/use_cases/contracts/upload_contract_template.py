from __future__ import annotations

import re
from uuid import UUID

from core.domain.contract_template import ContractTemplate
from core.ports.outbound.contract_template_repository import ContractTemplateRepositoryPort
from core.ports.outbound.storage_port import StoragePort


def _extract_variables(docx_bytes: bytes) -> list[str]:
    """Extract {variable_name} placeholders from DOCX using python-docx."""
    try:
        import io
        from docx import Document

        doc = Document(io.BytesIO(docx_bytes))
        text_parts: list[str] = []

        for para in doc.paragraphs:
            text_parts.append(para.text)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text_parts.append(para.text)

        full_text = "\n".join(text_parts)
        pattern = re.compile(r'\{([a-zA-Z_][a-zA-Z0-9_]*)\}')
        found = pattern.findall(full_text)
        # deduplicate preserving order
        seen: set[str] = set()
        result: list[str] = []
        for v in found:
            if v not in seen:
                seen.add(v)
                result.append(v)
        return result
    except Exception:
        return []


class UploadContractTemplateUseCase:

    def __init__(
        self,
        template_repo: ContractTemplateRepositoryPort,
        storage: StoragePort,
    ) -> None:
        self._repo = template_repo
        self._storage = storage

    async def execute(
        self,
        tenant_id: UUID,
        name: str,
        description: str,
        docx_bytes: bytes,
    ) -> ContractTemplate:
        variables = _extract_variables(docx_bytes)

        template = ContractTemplate.create(
            tenant_id=tenant_id,
            name=name,
            description=description,
            file_key="",
            variables=variables,
        )
        template.file_key = f"contract-templates/{tenant_id}/{template.id}.docx"

        await self._storage.upload(
            key=template.file_key,
            data=docx_bytes,
            content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        )
        await self._repo.save(template)
        return template
