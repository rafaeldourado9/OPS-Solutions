import io
import re

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from core.ports.outbound.docx_template_engine_port import DocxTemplateEnginePort

PLACEHOLDER_RE = re.compile(r"\{([a-zA-Z_][a-zA-Z0-9_]*)\}")

_W_P = qn("w:p")   # paragraph element
_W_T = qn("w:t")   # text element within a run


class DocxTemplateEngine(DocxTemplateEnginePort):

    def extract_text(self, docx_bytes: bytes) -> str:
        """
        Extract all plain text from the DOCX, one paragraph per line.

        Scans the same scope as extract_placeholders (body, tables, headers,
        footers, text boxes) so the LLM sees the full document.
        """
        doc = Document(io.BytesIO(docx_bytes))
        lines: list[str] = []

        for p_elem in doc._element.iter(_W_P):
            para_text = "".join(t.text or "" for t in p_elem.iter(_W_T)).strip()
            if para_text:
                lines.append(para_text)

        return "\n".join(lines)

    def extract_placeholders(self, docx_bytes: bytes) -> list[str]:
        """
        Extract all {placeholder} keys from the DOCX.

        Scans the entire XML tree (main body, tables, text boxes, headers,
        footers) so no placeholder is missed regardless of document structure.
        Within each paragraph, all <w:t> text nodes are concatenated before
        applying the regex, so broken runs (Word splitting {placeholder} across
        multiple XML runs) are handled correctly.
        """
        doc = Document(io.BytesIO(docx_bytes))
        found: set[str] = set()

        for p_elem in doc._element.iter(_W_P):
            # Concatenate all w:t text nodes in this paragraph
            para_text = "".join(t.text or "" for t in p_elem.iter(_W_T))
            for match in PLACEHOLDER_RE.finditer(para_text):
                found.add(match.group(1))

        return sorted(found)

    def fill_template(self, docx_bytes: bytes, context: dict[str, str]) -> bytes:
        """
        Replace all {placeholder} occurrences in the DOCX with values from context,
        preserving per-run formatting throughout the entire document.
        """
        doc = Document(io.BytesIO(docx_bytes))

        for p_elem in doc._element.iter(_W_P):
            para = Paragraph(p_elem, doc)
            self._replace_in_paragraph(para, context)

        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()

    @staticmethod
    def _replace_in_paragraph(para: Paragraph, context: dict[str, str]) -> None:
        """
        Replace placeholders in a paragraph while preserving per-run formatting.

        Word sometimes splits {placeholder} across multiple runs, e.g.:
            run1.text = "{nome"
            run2.text = "_cliente}"
        This method handles that by assigning replacement text to the first run
        that contained the opening brace — all other run formatting is untouched.
        """
        runs = para.runs
        if not runs:
            return

        full_text = "".join(r.text for r in runs)
        if "{" not in full_text:
            return
        if not PLACEHOLDER_RE.search(full_text):
            return

        # Build char → run index mapping for the original text
        char_to_run: list[int] = []
        for i, run in enumerate(runs):
            char_to_run.extend([i] * len(run.text))

        # Collect all replacements in order
        repl_positions: list[tuple[int, int, str]] = []
        for m in PLACEHOLDER_RE.finditer(full_text):
            replacement = context.get(m.group(1), m.group(0))
            repl_positions.append((m.start(), m.end(), replacement))

        if not repl_positions:
            return

        def replacer(m: re.Match) -> str:
            return context.get(m.group(1), m.group(0))

        new_full = PLACEHOLDER_RE.sub(replacer, full_text)
        if new_full == full_text:
            return

        # For each character in new_full, decide which run it belongs to.
        # Unchanged characters keep their original run assignment.
        # Replacement characters go to the first run that contained the opening brace.
        new_char_runs: list[int] = []
        repl_idx = 0
        orig_i = 0

        while orig_i < len(full_text):
            if repl_idx < len(repl_positions) and orig_i == repl_positions[repl_idx][0]:
                start, end, repl = repl_positions[repl_idx]
                first_run = char_to_run[start] if start < len(char_to_run) else len(runs) - 1
                new_char_runs.extend([first_run] * len(repl))
                orig_i = end
                repl_idx += 1
            else:
                run_idx = char_to_run[orig_i] if orig_i < len(char_to_run) else len(runs) - 1
                new_char_runs.append(run_idx)
                orig_i += 1

        # Accumulate text per run
        run_new_texts: list[str] = [""] * len(runs)
        for char, run_idx in zip(new_full, new_char_runs):
            run_new_texts[run_idx] += char

        # Write back — only .text is touched; rPr (bold/italic/color/size) is untouched
        for run, new_text in zip(runs, run_new_texts):
            run.text = new_text

    def inject_placeholders(
        self, docx_bytes: bytes, injections: dict[str, str]
    ) -> bytes:
        """
        Replace literal text spans with {placeholder_key} syntax.

        Works on "virgin" documents (no existing {placeholders}).
        Applies the same run-reassignment algorithm as fill_template so that
        run formatting (bold, color, size) is preserved after injection.

        injections: {"R$ 0,00": "valor_total", "NOME DO CLIENTE": "nome_cliente"}
        """
        if not injections:
            return docx_bytes

        # Sort by length descending to avoid partial-match collisions
        # (e.g. "R$ 1.000,00" must be checked before "1.000,00")
        sorted_items = sorted(injections.items(), key=lambda x: len(x[0]), reverse=True)

        # Build a single regex that matches any of the original texts
        pattern = re.compile(
            "|".join(re.escape(orig) for orig, _ in sorted_items)
        )
        # Map escaped text back to placeholder key
        text_to_key = {orig: key for orig, key in sorted_items}

        doc = Document(io.BytesIO(docx_bytes))

        for p_elem in doc._element.iter(_W_P):
            para = Paragraph(p_elem, doc)
            self._inject_in_paragraph(para, pattern, text_to_key)

        output = io.BytesIO()
        doc.save(output)
        return output.getvalue()

    @staticmethod
    def _inject_in_paragraph(
        para: Paragraph,
        pattern: re.Pattern,
        text_to_key: dict[str, str],
    ) -> None:
        """
        Replace literal spans in a paragraph with {placeholder_key} tokens,
        preserving per-run formatting using the same char→run mapping as fill_template.
        """
        runs = para.runs
        if not runs:
            return

        full_text = "".join(r.text for r in runs)
        if not pattern.search(full_text):
            return

        # Build char → run index mapping
        char_to_run: list[int] = []
        for i, run in enumerate(runs):
            char_to_run.extend([i] * len(run.text))

        # Collect replacement positions in order
        repl_positions: list[tuple[int, int, str]] = []
        for m in pattern.finditer(full_text):
            replacement = "{" + text_to_key[m.group(0)] + "}"
            repl_positions.append((m.start(), m.end(), replacement))

        if not repl_positions:
            return

        def replacer(m: re.Match) -> str:
            return "{" + text_to_key[m.group(0)] + "}"

        new_full = pattern.sub(replacer, full_text)
        if new_full == full_text:
            return

        # Re-map characters to runs
        new_char_runs: list[int] = []
        repl_idx = 0
        orig_i = 0

        while orig_i < len(full_text):
            if repl_idx < len(repl_positions) and orig_i == repl_positions[repl_idx][0]:
                start, end, repl = repl_positions[repl_idx]
                first_run = char_to_run[start] if start < len(char_to_run) else len(runs) - 1
                new_char_runs.extend([first_run] * len(repl))
                orig_i = end
                repl_idx += 1
            else:
                run_idx = char_to_run[orig_i] if orig_i < len(char_to_run) else len(runs) - 1
                new_char_runs.append(run_idx)
                orig_i += 1

        run_new_texts: list[str] = [""] * len(runs)
        for char, run_idx in zip(new_full, new_char_runs):
            run_new_texts[run_idx] += char

        for run, new_text in zip(runs, run_new_texts):
            run.text = new_text
